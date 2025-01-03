import google.generativeai as genai
from docx import Document
from datetime import datetime

# Configure the Gemini API
GEMINI_API_KEY = "AIzaSyDseWN7yM71gg3oHQ0nnvH9zaMpOYS3Ceo"
genai.configure(api_key=GEMINI_API_KEY)

# Define the model
model = genai.GenerativeModel("gemini-1.5-flash")

# Function to generate a script
def generate_script():
    prompt = (
        "Create a creative, attractive, and professional script for a 60-second YouTube short video. "
        "The video is targeted at people who love cooking italian pizza, baking, and exploring food recipes. "
        "The tone should be engaging and include a call-to-action for viewers to like, share, and subscribe."
        "Include only text without any coding signs or video length or astric signs only the speaking text or pargraph explenation, highlights vedio action explenation"
    )
    
    print("Sending prompt to the AI model...")
    response = model.generate_content(prompt)
    script_text = response.text

    # Save the script to a Word document
    date_str = datetime.now().strftime("%Y-%m-%d")
    file_name = f"Cooking_Script_{date_str}.docx"

    document = Document()
    #document.add_heading("YouTube Short Video Script", level=2)
    document.add_paragraph(script_text)
    document.save(file_name)

    print(f"The script has been successfully saved as '{file_name}'.")

# Run the script generation
if __name__ == "__main__":
    generate_script()
